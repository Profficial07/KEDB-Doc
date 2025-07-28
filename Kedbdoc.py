import streamlit as st
from datetime import date
from docx import Document
from docx.shared import Inches
import os
import re
from io import BytesIO

def main():
    st.set_page_config(page_title="KEDB Generator", layout="wide")
    st.title("KEDB Document")

    # Optional Sidebar
    st.sidebar.title("KEDB")
    st.sidebar.markdown("Fill in all required details to generate a structured KEDB document.")
    # st.sidebar.image("logo.png", width=200)  # Optional logo

    kedb_no = get_next_kedb_number()
    st.markdown(f"### 📄 KEDB Number: `{kedb_no}`")

    # ---- Page 1 - Issue Summary ----
    with st.expander("📝 Page 1 - Issue Summary", expanded=True):
        short_desc = st.text_input("Short Description")
        issue_type = st.selectbox("Issue Type", ["System", "Data", "Configuration", "Other"])

    # ---- Page 2 - Revision History ----
    with st.expander("📄 Page 2 - Document Revision History", expanded=True):
        today = date.today().strftime("%Y-%m-%d")
        creator = st.text_input("Creator Name")
        version = st.text_input("Version Number", "1.0")
        approver = st.text_input("Approver Name (Optional)")

    # ---- Page 3 - Detailed Description ----
    with st.expander("🔍 Page 3 - Detailed Description", expanded=True):
        detailed_desc = st.text_area("Detailed Description", height=200)

    # ---- Page 4 - Root Cause and Fix ----
    with st.expander("🛠️ Page 4 - Root Cause & Workaround", expanded=True):
        root_cause = st.text_area("Root Cause", height=150)
        root_images = st.file_uploader("Upload Snap(s) for Root Cause", accept_multiple_files=True, type=["png", "jpg", "jpeg"])
        
        fix = st.text_area("Workaround / Fix", height=150)
        fix_images = st.file_uploader("Upload Snap(s) for Workaround / Fix", accept_multiple_files=True, type=["png", "jpg", "jpeg"])

    # ---- Generate Button ----
    if st.button("🚀 Generate KEDB Document"):
        output_path = f"output_docs/{kedb_no}.docx"
        generate_doc(kedb_no, short_desc, issue_type, today, creator, version, approver,
                     detailed_desc, root_cause, root_images, fix, fix_images, output_path)

        # Download Button
        with open(output_path, "rb") as file:
            file_bytes = file.read()
            st.success(f"✅ Document `{kedb_no}.docx` created successfully!")
            st.download_button(
                label="📥 Download KEDB Document",
                data=file_bytes,
                file_name=f"{kedb_no}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ---- Document Generation ----
def generate_doc(kedb_no, short_desc, issue_type, today, creator, version, approver,
                 detailed_desc, root_cause, root_images, fix, fix_images, output_path):
    doc = Document()

    doc.add_heading(f"KEDB Entry - {kedb_no}", 0)

    # Page 1
    doc.add_heading("Issue Summary", level=1)
    doc.add_paragraph(f"KEDB Number: {kedb_no}")
    doc.add_paragraph(f"Short Description: {short_desc}")
    doc.add_paragraph(f"Issue Type: {issue_type}")
    doc.add_page_break()

    # Page 2 - Revision History with Table
    doc.add_heading("Document Revision History", level=1)

    # Create a table with 4 columns: Date, Version, Creator, Approver
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'  # Optional styling

    # Set column headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Version'
    hdr_cells[2].text = 'Creator'
    hdr_cells[3].text = 'Approver'

    # Add entry row
    row_cells = table.add_row().cells
    row_cells[0].text = today
    row_cells[1].text = version
    row_cells[2].text = creator
    row_cells[3].text = approver if approver else 'N/A'

    doc.add_page_break()


    # Page 3
    doc.add_heading("1.Detailed Description of Problem", level=1)
    doc.add_paragraph(detailed_desc)
    doc.add_page_break()

    # Page 4
    doc.add_heading("2.Root Cause", level=1)
    doc.add_paragraph(root_cause)
    if root_images:
        for img in root_images:
            doc.add_picture(img, width=Inches(5))
            doc.add_paragraph("")

    doc.add_heading("3.Workaround / Fix", level=1)
    doc.add_paragraph(fix)
    if fix_images:
        for img in fix_images:
            doc.add_picture(img, width=Inches(5))
            doc.add_paragraph("")

    os.makedirs("output_docs", exist_ok=True)
    doc.save(output_path)

# ---- Smart KEDB Number Logic ----
def get_next_kedb_number():
    output_dir = "output_docs"
    os.makedirs(output_dir, exist_ok=True)
    files = os.listdir(output_dir)

    kedb_numbers = []
    pattern = re.compile(r"KEDB(\d{4})\.docx")

    for f in files:
        match = pattern.match(f)
        if match:
            kedb_numbers.append(int(match.group(1)))

    if kedb_numbers:
        next_number = max(kedb_numbers) + 1
    else:
        next_number = 1

    return f"KEDB{next_number:04}"


if __name__ == "__main__":
    main()
